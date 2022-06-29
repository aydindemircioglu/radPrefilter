#!/usr/bin/python3

from collections import OrderedDict
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from contextlib import contextmanager


import mlflow
from mlflow.tracking.client import MlflowClient

from sklearn.utils import resample
from sklearn.metrics import roc_curve, auc, roc_auc_score, precision_recall_curve

from glob import glob
from joblib import dump, load
from matplotlib import cm
from matplotlib import pyplot
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
from matplotlib.transforms import Bbox
from PIL import Image
from PIL import ImageDraw, ImageFont
from scipy.stats import spearmanr, pearsonr
from typing import Dict, Any
import copy
import cv2
import hashlib
import itertools
import json
import math
import matplotlib
import matplotlib.colors as colors
import matplotlib.pyplot as plt
import multiprocessing
import numpy as np
import os
import pandas as pd
import pathlib
import pickle
import pylab
import random
import scipy.cluster.hierarchy as sch
import seaborn as sns
import shutil
import sys
import tempfile
import time

# delong
import rpy2
import rpy2.robjects as robjects
from rpy2.robjects.packages import importr
from rpy2.robjects import pandas2ri
pandas2ri.activate()
pROC = importr('pROC')


#from utils import *
from nestedExperiment import preprocessData, selectFeatureSubset
from helpers import *
from parameters import *
#from evaluate_utils import *

cFigNumber = 1
document = None



def countData (dList):
    # has no diagnosis, but target and patientID
    for d in dList:
        X = loadData (d, drop = False)
        print ("Shape of ", d, str(len(set(X["Patient"]))))

    #X = X.drop(["Diagnosis", "Patient"], axis = 1)
    #print ("\tLoaded data with shape", X.shape)
    # dump the features while we are at it
    if d == "Melanoma":
        document = Document()
        font = document.styles['Normal'].font
        font.name = 'Arial'
        document.add_heading('Extracted Features\n', level = 1)
        fList = {}
        for f in filterNames:
            fList[f] = set(["_".join(k.split("_")[1:]) for k in X.keys() if f in k])
        commonFeats = set.intersection(*(fList[f] for f in fList.keys()))
        document.add_heading("Common features for all preprocessing filters:", level = 3)
        fTxt = ''
        for c in sorted(list(commonFeats)):
            fTxt = fTxt + c + "\n"
        document.add_paragraph(fTxt)

        for f in filterNames:
            feats = []
            for k in [k for k in X.keys() if f in k]:
                b = "_".join(k.split("_")[1:])
                if b in commonFeats:
                    continue
                feats.append(b)
            if len(feats) > 0:
                fTxt = "Features in " + getName(f)
                document.add_heading(fTxt, level = 3)
                fTxt = ''
                for k in feats:
                    fTxt = fTxt + str(k) + "\n"
                document.add_paragraph(fTxt)
        document.save('./paper/FeatureNames.docx')

    return X.copy()


def getStr(k):
    if k == "best":
        return "Tuned"
    if k == "org":
        return "Original"
    if k == "all":
        return "All"
    raise Exception ("What "+ k)


def getName (s):
    if s == "all":
        return "All"
    if s == "exponential_":
        return "Exponential"
    if s == "gradient_":
        return "Gradient"
    if s == "lbp-3D":
        return "LBP"
    if s == "log-sigma-":
        return "LoG"
    if s == "square_":
        return "Square"

    if s == "squareroot_":
        return "Squareroot"
    if s == "wavelet-":
        return "Wavelet"
    if s == "original_":
        return "Original"
    if s == "logarithm_":
        return "Logarithm"
    return (s)


def getNameList (s):
    return [getName(k) for k in s]


def  testAUC (Y, scoresA, scoresB, verbose = True):
        bVals = []
        t, pA, pB = np.asarray(Y), np.asarray(scoresA), np.asarray(scoresB)
        for b in range(2000):
            np.random.seed(b)
            idx = np.random.choice(len(t), len(t)) #list(range(len(t))), replace = True)
            b_pA = pA[idx]
            b_pB = pB[idx]
            b_t = t[idx]

            b_fprA, b_tprA, b_thresholdsA = roc_curve (b_t, b_pA)
            b_fprB, b_tprB, b_thresholdsB = roc_curve (b_t, b_pB)
            area_under_curveA = auc (b_fprA, b_tprA)
            area_under_curveB = auc (b_fprB, b_tprB)
            bVals.append(area_under_curveA - area_under_curveB)

        fprA, tprA, thresholdsA = roc_curve (Y, scoresA)
        area_under_curveA = auc (fprA, tprA)
        fprB, tprB, thresholdsB = roc_curve (Y, scoresB)
        area_under_curveB = auc (fprB, tprB)

        z = np.sum(area_under_curveA -area_under_curveB)/np.std(bVals)
        p =2*scipy.stats.norm.cdf(-np.abs(z))
        return (p, area_under_curveA, area_under_curveB)


def getResults (dList):
    mlflow.set_tracking_uri(trackingPath)
    if os.path.exists("./results/results.feather") == False:
        results = []
        for d in dList:
            current_experiment = dict(mlflow.get_experiment_by_name(d + "_final"))
            experiment_id = current_experiment['experiment_id']
            runs = MlflowClient().search_runs(experiment_ids=experiment_id, max_results=50000)
            len(runs)

            for r in runs:
                row = r.data.metrics
                row["UUID"] = r.info.run_uuid
                row["Model"] = r.data.tags["Version"]
                row["Parameter"] = r.data.tags["pID"]

                row["Parameter"] = row["Parameter"]
                row["Model"] = row["Model"]

                row["FSel"], row["Clf"] = row["Model"].split("_")
                row["Dataset"] = d

                row["nFeatures"] = eval(row["Parameter"])[row["FSel"]]["nFeatures"]
                row["featureSet"] = eval(row["Parameter"])[row["FSel"]]["featureSet"]

                row["Path"] = os.path.join(trackingPath,  str(experiment_id), str(r.info.run_uuid), "artifacts")
                results.append(row)

                # read timings
                apath = os.path.join(row["Path"], "timings.json")
                with open(apath) as f:
                    expData = json.load(f)
                row.update(expData)

                # read AUCs
                apath = os.path.join(row["Path"], "aucStats.json")
                with open(apath) as f:
                    aucData = json.load(f)
                row.update(aucData)

        print ("Pickling results")
        pickle.dump (results, open("./results/results.feather","wb"))
    else:
        print ("Restoring results")
        results = pickle.load(open("./results/results.feather", "rb"))

    return results


def getCI (predsX):
    Y = predsX["y_true"].values
    scoresA = predsX["y_pred"].values
    lower, auc, upper = pROC.ci(Y, scoresA, direction = "<")
    return lower, auc, upper


def delongTest (predsX, predsY):
    Y = predsX["y_true"].values
    scoresA = predsX["y_pred"].values
    scoresB = predsY["y_pred"].values
    rocA = pROC.roc (Y, scoresA, direction = "<")
    rocB = pROC.roc (Y, scoresB, direction = "<")

    aucA = pROC.auc(Y, scoresA)
    aucB = pROC.auc(Y, scoresB)
    #print ("AUC A:" + str(aucA))
    #print ("AUC B:" + str(aucB))
    robjects.globalenv['rocA'] = rocA
    robjects.globalenv['rocB'] = rocB

    z = rpy2.robjects.packages.reval ("library(pROC);z = roc.test(rocA, rocB, method= 'delong', progress='none'); p = z$p.value")
    z = robjects.r.z
    p = robjects.r.p[0]
    return p, aucA, aucB



def addText (finalImage, text = '', org = (0,0), fontFace = '', fontSize = 12, color = (255,255,255)):
     # Convert the image to RGB (OpenCV uses BGR)
     #tmpImg = cv2.cvtColor(finalImage, cv2.COLOR_BGR2RGB)
     tmpImg = finalImage
     pil_im = Image.fromarray(tmpImg)
     draw = ImageDraw.Draw(pil_im)
     font = ImageFont.truetype(fontFace + ".ttf", fontSize)
     draw.text(org, text, font=font, fill = color)
     #tmpImg = cv2.cvtColor(np.array(pil_im), cv2.COLOR_RGB2BGR)
     tmpImg = np.array(pil_im)
     return (tmpImg.copy())



def addBorder (img, pos, thickness):
    if pos == "H":
        img = np.hstack([255*np.ones(( img.shape[0],int(img.shape[1]*thickness), 3), dtype = np.uint8),img])
    if pos == "V":
        img = np.vstack([255*np.ones(( int(img.shape[0]*thickness), img.shape[1], 3), dtype = np.uint8),img])
    return img



def evalTimes (dList, results):
    fTimes = {}
    for d in dList:
        fTimes[d] = {}
        for k in range(10):
            expName = d + "_" + str(k)
            results = pickle.load(open("./results/"+str(expName)+".feather", "rb"))
            aTable = results.query("Dataset == @d").copy()
            aTable["Time"] = np.sum(aTable[[k for k in aTable.keys() if "Time" in k]], axis = 1)
            fsets = sorted(list(set(results["featureSet"].values)))
            for f in fsets:
                fTimes [d][f] = np.sum(aTable.query("featureSet == @f ")["Time"])

    # add retraining times
    for d in dList:
        aTable = results.query("Dataset == @d").copy()
        aTable["Time"] = np.sum(aTable[[k for k in aTable.keys() if "Time" in k]], axis = 1)
        fsets = sorted(list(set(results["featureSet"].values)))
        for f in fsets:
            fTimes [d][f] = fTimes[d][f] + np.sum(aTable.query("featureSet == @f ")["Time"])
    tf = pd.DataFrame(fTimes)
    timepfs = tf.sum(axis = 1)
    print ("Absolute computation time:  tuning = {0}".format( (np.sum(timepfs) )))
    print ("Absolute computation time:  original = {0}".format( timepfs["original_"]))
    print ("Absolute computation time:  all = {0}".format( timepfs["all"]))
    print ("Increase in computation time:  tuning vs all = {0}".format( (np.sum(timepfs) - timepfs["all"])/timepfs["all"]))
    print ("Increase in computation time: tuning vs original = {0}".format( (np.sum(timepfs) - timepfs["original_"])/timepfs["original_"]))
    print ("Increase in computation time: original vs all = {0}".format( (timepfs["all"]/timepfs["original_"])))
    pass


# cross similairty of a pattern
# if one features is constant, we set corr() = 1
# if not, we get corr(lbp,lbp) < 1...
def getCorr (fA, fB):
    cu = []
    for u in range(fA.shape[1]):
        cv = []
        for v in range(fB.shape[1]):
            if len(set(fA.values[:,u])) == 1 or len(set(fB.values[:,v])) == 1:
                corr = 1.0
            else:
                corr, pval = pearsonr (fA.values[:,u], fB.values[:,v])
            cv.append(np.abs(corr))
        cu.append(np.max(cv))
    CS = np.mean(cu)
    return CS



def get_correlation (results):
    if os.path.exists("./results/corr.feather") == False:
        # for each dataset
        corrs = {}
        for d in dList:
            # get data
            data = loadData(d)
            y = data["Target"]
            X = data.drop(["Target"], axis = 1)
            X, y = preprocessData (X, y)

            # compute correlation between the two fsets
            dsets = {}
            for fset in filterNames:
                # only take those features we need
                fExp = ([None,{"featureSet": fset}],)
                dsets[fset] = list(selectFeatureSubset (X, y, fExp))

            # remove original from all but original
            oKeys = list(dsets["original_"][0].keys())
            for fset in filterNames:
                if fset == "original_":
                    continue
                tmpX = dsets[fset][0].drop(oKeys, axis = 1)
                dsets[fset][0] = tmpX.copy()

            # for fset in filterNames:
            #     print(dsets[fset][0].shape)

            rMat = np.zeros( (len(filterNames), len(filterNames) ) )
            rMat = pd.DataFrame(rMat, index = filterNames, columns = filterNames)
            for fsetA in filterNames:
                for fsetB in filterNames:
                    rMat.at[fsetA, fsetB] = getCorr (dsets[fsetA][0], dsets[fsetB][0])
            corrs[d] = rMat
        print ("Pickling results")
        pickle.dump (corrs, open("./results/corr.feather","wb"))
    else:
        print ("Restoring results")
        corrs = pickle.load(open("./results/corr.feather", "rb"))
    return corrs



def plot_correlations (corrs):
    Z = np.mean([corrs[k] for k in corrs], axis = 0)
    Z = pd.DataFrame(Z, columns = corrs["Melanoma"].columns, index = corrs["Melanoma"].columns)
    Z = (Z + Z.T)/2

    # reorder, because filternames is messed up
    newOrder = ["original_"] + sorted(filterNames[1:])
    Z = Z[newOrder].loc[newOrder].copy()

    rMat = Z
    rMat = rMat.round(3)
    scMat = rMat.copy()
    strMat = rMat.copy()
    strMat = strMat.astype( dtype = "str")

    # this computes the mean correlation between the feature sets.
    # but the correlation of original to each is read off the matrix
    np.mean(Z,axis=1).sort_values()

    if 1 == 1:
        plt.rc('text', usetex=True)
        plt.rcParams.update({
            "font.family": "sans-serif",
            "font.sans-serif": ["Arial"]})
        plt.rcParams['text.usetex'] = True
        plt.rcParams['text.latex.preamble'] = r'''
            \usepackage{mathtools}
            \usepackage{helvet}
            \renewcommand{\familydefault}{\sfdefault}        '''

        fig, ax = plt.subplots(figsize = (13,14), dpi = DPI)
        sns.set(style='white')
        #ax = sns.heatmap(scMat, annot = cMat, cmap = "Blues", fmt = '', annot_kws={"fontsize":21}, linewidth = 2.0, linecolor = "black")
        dx = np.asarray(scMat, dtype = np.float64)
        pal = sns.light_palette("#8888bb", reverse=False, as_cmap=True)
        pal = "rocket"
        tnorm = colors.TwoSlopeNorm(vmin=0.0, vcenter=0.5, vmax=1.0)
        ax.imshow(dx, cmap=pal, norm = tnorm, interpolation='nearest', aspect = 0.49)

        # Major ticks
        mh, mw = scMat.shape
        ax.set_xticks(np.arange(0, mw, 1))
        ax.set_yticks(np.arange(0, mh, 1))

        # Minor ticks
        ax.set_xticks(np.arange(-.5, mw, 1), minor=True)
        ax.set_yticks(np.arange(-.5, mh, 1), minor=True)

        # Gridlines based on minor ticks
        ax.grid(which='minor', color='black', linestyle='-', linewidth=2)

        for i, c in enumerate(scMat.index):
            for j, f in enumerate(scMat.keys()):
                    ax.text(j, i, strMat.at[c, f],    ha="center", va="center", color="k", fontsize = 22)
        plt.tight_layout()

        ax.set_xticklabels(getNameList(rMat.keys()), rotation = 45, ha = "right", fontsize = 22)
        ax.set_yticklabels(getNameList(rMat.index), rotation = 0, ha = "right", fontsize = 22)
        ax.yaxis.set_tick_params ( labelsize= 22)
        fig.savefig("./paper/Figure_4.png", facecolor = 'w', bbox_inches='tight')
        plt.close('all')

    plt.rc('text', usetex=False)
    plt.close('all')
    pass


def getPreds (dList, results):
    # assemble preds first
    allPreds = {}
    for d in dList:
        allPreds[d] = {}
        for fset in set(results["featureSet"]):
            fsTbl = results.query("Dataset == @d and featureSet == @fset")

            # create auc curve here
            preds = []
            for (i, (idx, row)) in enumerate(fsTbl.iterrows()):
                preds.append(pd.read_csv(os.path.join(row["Path"], "preds.csv")))
            allPreds[d][fset] =  pd.concat(preds).sort_values(["Unnamed: 0"]).reset_index(drop = True)


    aucTblR = {}
    for d in dList:
        aucTblR[d] = {}
        for fset in set(results["featureSet"]):
            _, aucA, _ = getCI (allPreds[d][fset])
            aucTblR[d][fset] = aucA

    tR = pd.DataFrame(aucTblR).T
    tR["max"] = np.max(tR,axis=1)
    return allPreds, tR


# je groesser der datensatz, desto weniger bringt optimieren des preprocess etwas verglichen mit all
# je kleiner der datensatz, desto mehr bringt optimieren des preprocess etwas verglichen mit all

# je kleiner der datensatz, desto weniger bringt optimieren des preprocess etwas verglichen mit original
# je groesser der datensatz, desto mehr bringt optimieren des preprocess etwas verglichen mit original

def createAUCplots (dList, results, tR, allPreds):
    # create AUC plots
    for d in dList:
        # paint only best, all, original-- all includes best!
        bestFilter = tR[filterNames + ["all"]].idxmax( axis = 1)[d]

        preds = {"best": allPreds[d][bestFilter], "org": allPreds[d]["original_"], "all": allPreds[d]["all"]}

        base_fpr = np.linspace(0, 1, 101)
        tprs = {}
        aucs = {}
        sens = {}
        spec = {}
        for k in ["best", "org", "all"]:
            tprs[k] = []
            aucs[k] = []
            sens[k] = []
            spec[k] = []

        for b in range(2000):
            np.random.seed(b)
            random.seed(b)

            # bootstrap identifcal
            idx = resample(range(len(preds["best"]["y_true"])))
            for k in ["best", "org", "all"]:
                bs_tA, bs_pA = [preds[k]["y_true"][idx], preds[k]["y_pred"][idx]]

                # python style
                p_fprA, p_tprA, p_thresholdsA = roc_curve (bs_tA, bs_pA)
                #area_under_curveA = auc (fprA, tprA)

                # R style
                oz = pROC.roc (np.array(bs_tA), np.array(bs_pA), direction="<")

                z = dict(oz.items())
                tprA = z["sensitivities"][::-1]
                fprA = (1-z["specificities"])[::-1]
                area_under_curveA = pROC.auc(oz)
                thresholdsA = z["thresholds"][::-1]

                aucs[k].append(area_under_curveA)
                tpr = np.interp(base_fpr, fprA, tprA)
                tpr[0] = 0.0
                tprs[k].append(tpr)
                csens, cspec = findOptimalCutoff (fprA, tprA, thresholdsA)
                sens[k].append(csens)
                spec[k].append(cspec)

        # base curve
        f, ax = plt.subplots(figsize = (7,7), dpi = DPI)
        colors = {"best": "#f7022a", "org": "#062af3", "all": "#028f1e"}
        colors = {"best": "r", "org": "g", "all": "b"}
        for k in ["best", "org", "all"]:
            # python style
            #fprA, tprA, thresholdsA = roc_curve (preds[k]["y_true"], preds[k]["y_pred"])
            #area_under_curveA = auc (fprA, tprA)

            # R style
            oz = pROC.roc (np.array(preds[k]["y_true"]), np.array(preds[k]["y_pred"]), direction = "<")
            z = dict(oz.items())
            tprA = z["sensitivities"][::-1]
            fprA = (1-z["specificities"])[::-1]
            area_under_curveA = pROC.auc(oz)[0]
            thresholdsA = z["thresholds"][::-1]


            # GET CI TEST
            _, auc_ci, _ = pROC.ci(np.array(preds[k]["y_true"]), np.array(preds[k]["y_pred"]), direction = "<")
            if np.abs(auc_ci - area_under_curveA) > 0.01:
                print (area_under_curveA, auc_ci)
                raise Exception ("H")

            csens, cspec = findOptimalCutoff (fprA, tprA, thresholdsA)
            print ("ORG", k, "AUC/Sens/Spec:", area_under_curveA, csens, cspec)
            print ("BOOT", k, "AUC/Sens/Spec:", np.mean(aucs[k]), np.mean(sens[k]), np.mean(spec[k]))

            tprsA = np.array(tprs[k])
            mean_tprsA = tprsA.mean(axis=0)
            stdA = tprsA.std(axis=0)
            tprsA_upper = np.minimum(mean_tprsA + stdA, 1)
            tprsA_lower = np.maximum(mean_tprsA - stdA, 0)

            plt.plot(base_fpr, mean_tprsA, colors[k], label =  'AUC  (' + getStr(k) + '): {0:0.2f}'.format(area_under_curveA) )
            plt.fill_between(base_fpr, tprsA_lower, tprsA_upper, color=colors[k], alpha=0.15)

        plt.xticks(fontsize=18)
        plt.yticks(fontsize=18)
        plt.xlabel('1 - Specificity', fontsize = 22, labelpad = 12)
        plt.ylabel('Sensitivity', fontsize= 22, labelpad = 12)

        ax.plot([0, 1], [0, 1],'k--')
        ax.set_xlim([-0.01, 1.01])
        ax.set_ylim([-0.01, 1.01])
        ax.set_aspect('equal', 'datalim')

        ax.set_title("AUC-ROC for " + d, fontsize = 26)
        ax.legend(loc="lower right", fontsize = 18)
        f.savefig("./results/AUCROC_" + d + ".png", facecolor = 'w')





# just read both figures and merge them
def join_AUC_plots():
    fontFace = "Arial"

    imA = cv2.imread("./results/AUCROC_CRLM.png")
    imA = addText (imA, "a", (40,imA.shape[0]-170), fontFace, 112, color= (0,0,0))

    imB = cv2.imread("./results/AUCROC_Desmoid.png")
    imB = addText (imB, "b", (40,imB.shape[0]-170), fontFace, 112, color=(0,0,0))

    imC = cv2.imread("./results/AUCROC_GIST.png")
    imC = addText (imC, "c", (40,imC.shape[0]-170), fontFace, 112, color=(0,0,0))

    imB = addBorder (imB, "H", 0.075)
    imC = addBorder (imC, "H", 0.075)
    imgU = np.hstack([imA, imB, imC])


    imA = cv2.imread("./results/AUCROC_HN.png")
    imA = addText (imA, "d", (40,imA.shape[0]-170), fontFace, 112, color= (0,0,0))

    imB = cv2.imread("./results/AUCROC_Lipo.png")
    imB = addText (imB, "e", (40,imB.shape[0]-170), fontFace, 112, color=(0,0,0))

    imC = cv2.imread("./results/AUCROC_Liver.png")
    imC = addText (imC, "f", (40,imC.shape[0]-170), fontFace, 112, color=(0,0,0))

    imB = addBorder (imB, "H", 0.075)
    imC = addBorder (imC, "H", 0.075)
    imgM = np.hstack([imA, imB, imC])


    imA = cv2.imread("./results/AUCROC_Melanoma.png")
    imA = addText (imA, "g", (40,imA.shape[0]-170), fontFace, 112, color= (0,0,0))

    imB = imB*0 + 255
    imC = imC*0 + 255
    imgL = np.hstack([imA, imB, imC])

    imgM = addBorder (imgM, "V", 0.075)
    imgL = addBorder (imgL, "V", 0.075)
    img = np.vstack([imgU, imgM, imgL])
    cv2.imwrite("./paper/Figure_2.png", img)



def plot_TradeOffs (dList, results, tR):
    def doPlot(df, d, v1, v2):
        for ctype in [v2]:
            spfList = df[[v1, ctype]]
            R, pval = pearsonr(*zip (*spfList.values))
            R2 = R*R
            print (R, pval)

            # fSels = [z[0] for z in spfList.index]
            # dSets = [z[1] for z in spfList.index]

            x, y = zip(*spfList.values)
            p, cov = np.polyfit(x, y, 1, cov=True)                     # parameters and covariance from of the fit of 1-D polynom.
            y_model = equation(p, x)                                   # model using the fit parameters; NOTE: parameters here are coefficients

            # Statistics
            n =  len(x)                                          # number of observations
            ps = p.size                                                 # number of parameters
            dof = n - ps                                                # degrees of freedom
            t = stats.t.ppf(0.975, n - ps)                              # used for CI and PI bands

            # Estimates of Error in Data/Model
            resid = y - y_model
            chi2 = np.sum((resid / y_model)**2)                        # chi-squared; estimates error in data
            chi2_red = chi2 / dof                                      # reduced chi-squared; measures goodness of fit
            s_err = np.sqrt(np.sum(resid**2) / dof)                    # standard deviation of the error


            # plot
            if 1 == 1:
                fig, ax = plt.subplots(figsize = (7, 7), dpi = DPI)
                # sns.scatterplot (x = x,y = y,  ax = ax)
                sns.scatterplot (x = v1, y = ctype,  data=df,  ax = ax, s = 50, color = ".0")
                ax.plot(x, y_model, "-", color="0.1", linewidth=1.5, alpha=1.0, label="Fit")

                x2 = np.linspace(np.min(x), np.max(x), 100)
                y2 = equation(p, x2)

                # Confidence Interval (select one)
                plot_ci_manual(t, s_err, n, x, x2, y2, ax=ax)
                #plot_ci_bootstrap(x, y, resid, ax=ax)

                # Prediction Interval
                pi = t * s_err * np.sqrt(1 + 1/n + (x2 - np.mean(x))**2 / np.sum((x - np.mean(x))**2))
                ax.fill_between(x2, y2 + pi, y2 - pi, color="None", linestyle="--")
                ax.spines["top"].set_color("0.5")
                ax.spines["bottom"].set_color("0.5")
                ax.spines["left"].set_color("0.5")
                ax.spines["right"].set_color("0.5")
                ax.get_xaxis().set_tick_params(direction="out")
                ax.get_yaxis().set_tick_params(direction="out")
                ax.xaxis.tick_bottom()
                ax.yaxis.tick_left()
                #ax.invert_xaxis()

                # Labels
                if d == "all_vs_org":
                    plt.title("Trade-off Original vs All", fontsize="24")
                if d == "all_vs_max":
                    plt.title("Trade-off All vs Best", fontsize="24")
                if d == "org_vs_max":
                    plt.title("Trade-off Original vs Best", fontsize="24")

                plt.xticks(fontsize=20)
                plt.yticks(fontsize=20)
                #if d == "all_vs_org":
                plt.ylabel('Difference in AUC-ROC', fontsize = 22, labelpad = 12)

                plt.xlabel('Sample Size', fontsize = 22, labelpad = 12)
                #ax.set_xticks([50,250,500,750])

                right = 0.95
                ypos = 0.07 #0.93s
                legtext = ''
                if len(legtext ) > 0:
                    ypos = 0.07
                    legtext=legtext+"\n"

                plt.rcParams.update({
                    "text.usetex": True,
                    "font.family": "Arial",
                    "font.sans-serif": ["Arial"]})
                legpost = ''
                bbox_props = dict(fc="w", ec="0.5", alpha=0.9)
                pTxt = (' = {0:0.2f} ($p$ = {1:0.3f})').format(R2, pval)
                plt.text (right, ypos,
                          (legtext +  "$R^2$" + pTxt),
                          horizontalalignment='right',
                          size = 24, bbox  = bbox_props,
                          transform = ax.transAxes)
                plt.rcParams.update({
                    "text.usetex": False,
                    "font.family": "Arial",
                    "font.sans-serif": ["Arial"]})

                #ax.set_title("Stability vs AUC (" + d + ")", fontsize = 28)
                print ("Bias for", d)

                fig.tight_layout()
                fig.savefig("./results/Tradeoff_" + str(d) + ".png", facecolor = 'w')


    D = pd.DataFrame(tR["all"] - tR["original_"])
    for (i, (idx, row)) in enumerate(D.iterrows()):
        N = len(allPreds[idx]["all"])
        D.at[idx, "N"]  =N
    doPlot (D, d = "all_vs_org", v1 = "N", v2 = 0)

    D = pd.DataFrame(tR["all"] - tR["max"])
    for (i, (idx, row)) in enumerate(D.iterrows()):
        D.at[idx, "N"]  = len(allPreds[idx]["all"])
    doPlot (D, d = "all_vs_max", v1 = "N", v2 = 0)

    D = pd.DataFrame(tR["original_"] - tR["max"])
    for (i, (idx, row)) in enumerate(D.iterrows()):
        D.at[idx, "N"]  = len(allPreds[idx]["all"])
    doPlot (D, d = "org_vs_max", v1 = "N", v2 = 0)

    plt.close('all')
    pass


# just read both figures and merge them
def join_tradeoff_plots():
    fontFace = "Arial"

    imA = cv2.imread("./results/Tradeoff_all_vs_org.png")
    imA = addText (imA, "a", (40,imA.shape[0]-170), fontFace, 112, color= (0,0,0))

    imB = cv2.imread("./results/Tradeoff_org_vs_max.png")
    imB = addText (imB, "b", (40,imA.shape[0]-170), fontFace, 128, color=(0,0,0))

    imC = cv2.imread("./results/Tradeoff_all_vs_max.png")
    imC = addText (imC, "c", (40,imA.shape[0]-170), fontFace, 112, color=(0,0,0))

    imB = addBorder (imB, "H", 0.075)
    imC = addBorder (imC, "H", 0.075)
    img = np.hstack([imA, imB, imC])
    #Image.fromarray(img[::6,::6,:])

    cv2.imwrite("./paper/Figure_3.png", img)


def createTable (dList, results):
    rocTbl = {}
    allPreds = {}
    allAUCs = {}
    for d in dList:
        allPreds[d] = {}
        allAUCs[d] = {}
        for fset in set(results["featureSet"]):
            fsTbl = results.query("Dataset == @d and featureSet == @fset")

            # create auc curve here
            preds = []
            for (i, (idx, row)) in enumerate(fsTbl.iterrows()):
                preds.append(pd.read_csv(os.path.join(row["Path"], "preds.csv")))

            allPreds[d][fset] =  pd.concat(preds).sort_values(["Unnamed: 0"]).reset_index(drop = True)
            allAUCs[d][fset] = np.round(getCI(allPreds[d][fset]),2)
        allAUCs[d]
        pAO = np.round(delongTest (allPreds[d]["all"], allPreds[d]["original_"])[0], 3)
        bestFilter = pd.DataFrame(allAUCs[d]).idxmax(axis = 1).values[0]
        pBO  = np.round(delongTest (allPreds[d][bestFilter], allPreds[d]["original_"])[0], 3)
        pBA = np.round(delongTest (allPreds[d][bestFilter], allPreds[d]["all"])[0], 3)

        oStr = str(allAUCs[d]["original_"][1]) +  " ("+str(allAUCs[d]["original_"][0]) + "-" + str(allAUCs[d]["original_"][2]) + ")"
        bStr = str(allAUCs[d][bestFilter][1]) +  " ("+str(allAUCs[d][bestFilter][0]) + "-" + str(allAUCs[d][bestFilter][2]) + ")"
        aStr = str(allAUCs[d]["all"][1]) +  " ("+str(allAUCs[d]["all"][0]) + "-" + str(allAUCs[d]["all"][2]) + ")"

        rocTbl[d] = {"Original": oStr, "All":  aStr, "Tuned": bStr,  \
                            "Delta AO": allAUCs[d]["all"][1] - allAUCs[d]["original_"][1], "pAO":  pAO, \
                            "Delta BO": allAUCs[d][bestFilter][1] - allAUCs[d]["original_"][1], "pBO":  pBO, \
                            "Delta BA": allAUCs[d][bestFilter][1] - allAUCs[d]["all"][1], "pBA":  pBA, \
                            "Best Filter": bestFilter }


    rocTblDF = pd.DataFrame(rocTbl).T
    rocTblDF.to_excel("./paper/Table_3.xlsx")
    return rocTbl


if __name__ == "__main__":
    print ("Hi.")

    # obtain results
    print ("Generating results")
    results = getResults (dList)

    countData(dList)

    corrs = get_correlation (results)
    plot_correlations (corrs)

    evalTimes (dList, results)
    allPreds, tR = getPreds (dList, results)

    createAUCplots (dList, results, tR, allPreds)
    join_AUC_plots()

    rocTbl = createTable  (dList, results)

    # tradeoffs
    plot_TradeOffs(dList, results, tR)
    join_tradeoff_plots()



#
